import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.table import _Cell

import config as cfg
import caching
import tier_parsing
from utils import docx_funcs


class DocumentCreator:
    def __init__(
            self,
            grades: pd.DataFrame,
            left_margin: float = 3.18,
            right_margin: float = 3.18,
            top_margin: float = 2.54,
            bottom_margin: float = 2.54
    ):
        self.card_count: int = 0
        self.grades: pd.DataFrame = grades

        self.left_margin: float = left_margin
        self.right_margin: float = right_margin
        self.top_margin: float = top_margin
        self.bottom_margin: float = bottom_margin

        self.document: Document = Document()
        self.update_margins()

    def update_margins(self):
        docx_funcs.update_margins(
            self.document,
            top_margin=self.top_margin,
            bottom_margin=self.bottom_margin,
            left_margin=self.left_margin,
            right_margin=self.right_margin
        )

    def add_card_to_document(self, card_name: str):
        style = self.document.styles['Heading 3']
        docx_funcs.apply_default_font(style.font)

        heading = self.document.add_heading("", level=3)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = heading.add_run(card_name)
        card_data = caching.get_card_data(card_name)
        card_url = card_data["scryfall_uri"]

        docx_funcs.add_hyperlink(self.document, run, card_url)
        table = self.document.add_table(rows=3, cols=4)
        height, width = caching.download_card_image(card_name)
        docx_funcs.add_image_to_cell(table.cell(0, 0), height, width, cfg.TEMP_LOC)
        table.cell(0, 0).merge(table.cell(1, 2))

        self.add_grade(table.cell(0, 3), card_name, 'Marc')
        self.add_grade(table.cell(1, 3), card_name, 'Alex')

        paragraph = table.cell(2, 0).paragraphs[0]
        paragraph.style = 'List Bullet'
        table.cell(2, 0).merge(table.cell(2, 3))

        self.card_count += 1
        if self.card_count % 2 == 0:
            self.document.add_page_break()

    def add_grade(self, cell: _Cell, card_name: str, grader: str):
        grade = self.grades.loc[card_name][grader].strip()
        if grade.startswith('BA'):
            grade = f"{grade[2:]}, \nBuild-Around"
        if grade.startswith('SYN'):
            grade = f"{grade[3:]}, \nSynergy"

        cell.text = f"{grader}: {grade}"
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.runs[0]
        r.bold = True
        docx_funcs.apply_font(r.font, size=12, color=(0x00, 0x00, 0x00))

    def save_as(self, file_name: str):
        self.document.save(file_name)


def create_document(
        grades: pd.DataFrame,
        card_names: list[str],
        doc_name="test.docx",
        hori_margin: float = 3.18,
        vert_margin: float = 2.54
):
    doc_handler = DocumentCreator(
        grades,
        left_margin=hori_margin,
        right_margin=hori_margin,
        top_margin=vert_margin,
        bottom_margin=vert_margin
    )

    for card_name in card_names:
        card_name = card_name.strip()
        card_data = caching.get_card_data(card_name)

        if card_data:
            doc_handler.add_card_to_document(card_name)
        else:
            print(f"Card not found: {card_name}")

    doc_handler.save_as(doc_name)


def main(grades: pd.DataFrame, card_list: list[str] = None):
    if card_list is None:
        card_list = [card_name for card_name in grades.index]

    patched_cards = [caching.get_card_data(card_name)['name'] for card_name in card_list]
    create_document(grades, patched_cards, "test_3.docx", 1.5, 1)


if __name__ == "__main__":
    caching.populate_cache(['MOM', 'MUL'])
    set_grades = tier_parsing.parse_chord_excel(cfg.TIER_LIST_LOC, tier_parsing.patch_chord_excel)

    test_order = [
        "Botanical Brawler",
        "Invasion of Moag",
        "Tarkir Duneshaper",
    ]

    main(set_grades, test_order)
