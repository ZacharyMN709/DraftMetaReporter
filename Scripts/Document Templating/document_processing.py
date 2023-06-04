from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx.opc.constants
import docx.oxml.ns

import config as cfg
import caching
import tier_parsing
import image_processing


def set_cell_margins(cell, **kwargs):
    """
    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = docx.oxml.shared.OxmlElement('w:tcMar')

    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = docx.oxml.shared.OxmlElement("w:{}".format(m))
            node.set(docx.oxml.ns.qn('w:w'), str(kwargs.get(m)))
            node.set(docx.oxml.ns.qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def add_image_to_cell(cell, height: float, width: float, file_loc: str):
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run().add_picture(file_loc, height=Inches(height), width=Inches(width))


def apply_default_font(font):
    apply_font(font)


def apply_font(font, name: str = None, size: int = None, color: tuple[int, int, int] = None):
    font.name = name or cfg.FONT_NAME
    font.size = Pt(size or cfg.FONT_SIZE)
    font.color.rgb = RGBColor(*(color or cfg.FONT_COLOR))


def update_margins(document, top_margin: float, bottom_margin: float, left_margin: float, right_margin: float):
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(top_margin)
        section.bottom_margin = Cm(bottom_margin)
        section.left_margin = Cm(left_margin)
        section.right_margin = Cm(right_margin)


def add_hyperlink(document, run, url: str, is_external: bool = True):
    r = document.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=is_external)

    r_hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    r_hyperlink.set(docx.oxml.shared.qn('r:id'), r)
    r_hyperlink.set(docx.oxml.shared.qn('w:history'), '1')

    new_run = docx.oxml.shared.OxmlElement('w:r')
    r_properties = docx.oxml.shared.OxmlElement('w:rPr')

    new_run.append(r_properties)
    new_run.text = run.text
    r_hyperlink.append(new_run)

    run._r.getparent().replace(run._r, r_hyperlink)
    apply_default_font(run.font)


class DocumentCreator:
    def __init__(
        self,
        set_grades,
        left_margin=3.18,
        right_margin=3.18,
        top_margin=2.54,
        bottom_margin=2.54
    ):
        self.card_count = 0
        self.set_grades = set_grades

        self.left_margin = left_margin
        self.right_margin = right_margin
        self.top_margin = top_margin
        self.bottom_margin = bottom_margin

        self.document = Document()
        self.update_margins()

    def update_margins(self):
        update_margins(
            self.document,
            top_margin=self.top_margin,
            bottom_margin=self.bottom_margin,
            left_margin=self.left_margin,
            right_margin=self.right_margin
        )

    def add_card_to_document(self, card_name):
        style = self.document.styles['Heading 3']
        apply_default_font(style.font)

        heading = self.document.add_heading("", level=3)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = heading.add_run(card_name)
        card_data = caching.get_card_data(card_name)
        card_url = card_data["scryfall_uri"]

        add_hyperlink(self.document, run, card_url)
        table = self.document.add_table(rows=3, cols=4)
        height, width = image_processing.download_card_image(card_name)
        add_image_to_cell(table.cell(0, 0), height, width, cfg.TEMP_LOC)
        table.cell(0, 0).merge(table.cell(1, 2))

        self.add_grade(table.cell(0, 3), card_name, 'Marc')
        self.add_grade(table.cell(1, 3), card_name, 'Alex')

        paragraph = table.cell(2, 0).paragraphs[0]
        paragraph.style = 'List Bullet'
        table.cell(2, 0).merge(table.cell(2, 3))

        self.card_count += 1
        if self.card_count % 2 == 0:
            self.document.add_page_break()

    def add_grade(self, cell, card_name, grader):
        grade = self.set_grades.loc[card_name][grader].strip()
        if grade.startswith('BA'):
            grade = f"{grade[2:]}, \nBuild-Around"
        if grade.startswith('SYN'):
            grade = f"{grade[3:]}, \nSynergy"

        cell.text = f"{grader}: {grade}"
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.runs[0]
        r.bold = True
        apply_font(r.font, size=12, color=(0x00, 0x00, 0x00))

    def save_as(self, file_name):
        self.document.save(file_name)
